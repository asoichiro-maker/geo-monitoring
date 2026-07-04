"""
generate_monthly_report.py  (python-docx version, no Node.js required)
"""
from __future__ import annotations
try:
    from dotenv import load_dotenv; load_dotenv()
except ImportError:
    pass

import argparse, logging, os, subprocess, sys
from datetime import datetime, timezone
from pathlib import Path

logging.basicConfig(level=logging.INFO,
    format="%(asctime)s %(levelname)-8s %(message)s", datefmt="%H:%M:%S")
logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    logger.info("python-docx installing...")
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx",
                    "--break-system-packages"], check=True)
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

from mention_saver import MentionSaver

BASE_DIR          = Path(__file__).parent
REPORTS_DIR       = BASE_DIR / "reports"
DEFAULT_CLIENT_ID = "a1b2c3d4-0000-0000-0000-000000000001"

C_PRIMARY = RGBColor(0x1A, 0x3A, 0x5C)
C_ACCENT  = RGBColor(0xC9, 0xA8, 0x4C)
C_GRAY    = RGBColor(0x66, 0x66, 0x66)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT   = RGBColor(0xEA, 0xF0, 0xF6)
C_RED     = RGBColor(0xC0, 0x00, 0x00)
C_GREEN   = RGBColor(0x2E, 0x8B, 0x57)
C_LGRAY   = RGBColor(0xF5, 0xF5, 0xF5)

HX_PRIMARY = "1A3A5C"; HX_ACCENT = "C9A84C"; HX_LIGHT = "EAF0F6"
HX_WHITE   = "FFFFFF";  HX_LGRAY  = "F5F5F5"; HX_RED   = "C00000"
HX_PINK    = "FFF5F5";  HX_REDLT  = "FFCCCC"

LABELS = {"chatgpt":"ChatGPT","gemini":"Gemini",
          "perplexity":"Perplexity","ai_overview":"AI Overview"}
PROVIDERS = ["chatgpt","gemini","perplexity","ai_overview"]


def _bg(cell, hex_c):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"), hex_c); tcPr.append(shd)

def _borders(cell, hex_c="DDDDDD"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcB  = OxmlElement("w:tcBorders")
    for e in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{e}")
        el.set(qn("w:val"),"single"); el.set(qn("w:sz"),"4")
        el.set(qn("w:color"), hex_c); tcB.append(el)
    tcPr.append(tcB)

def _run(para, text, bold=False, size=11, color=None, font="Arial"):
    r = para.add_run(text)
    r.bold = bold; r.font.name = font; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return r

def _sp(para, bef=6, aft=6):
    para.paragraph_format.space_before = Pt(bef)
    para.paragraph_format.space_after  = Pt(aft)

def _h1(doc, text):
    p = doc.add_heading(level=1); p.clear(); _sp(p,14,6)
    r = p.add_run(text); r.font.name="Arial"; r.font.size=Pt(15)
    r.font.bold=True; r.font.color.rgb=C_PRIMARY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6")
    bot.set(qn("w:color"), HX_ACCENT); pBdr.append(bot); pPr.append(pBdr)

def _h2(doc, text):
    p = doc.add_heading(level=2); p.clear(); _sp(p,10,4)
    r = p.add_run(text); r.font.name="Arial"; r.font.size=Pt(12)
    r.font.bold=True; r.font.color.rgb=C_PRIMARY

def _hyperlink(para, url, display):
    r_id = para.part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hl = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), r_id)
    nr = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr"); rs = OxmlElement("w:rStyle")
    rs.set(qn("w:val"),"Hyperlink"); rPr.append(rs); nr.append(rPr)
    t = OxmlElement("w:t"); t.text = display; nr.append(t); hl.append(nr)
    para._p.append(hl)


def build_report(data: dict, out_path: Path) -> Path:
    year    = data["year"]; month = data["month"]
    mr      = data["mention_rate"]; mis = data["misinfo_rows"]
    sh_url  = data["spreadsheet_url"]
    period  = f"{year}年{month}月"
    total_q = sum(mr.get(p,{}).get("total",0)    for p in PROVIDERS)
    total_m = sum(mr.get(p,{}).get("mentioned",0) for p in PROVIDERS)
    overall = f"{total_m/total_q*100:.1f}%" if total_q else "0.0%"
    n_mis   = len(mis)

    doc = Document()
    for sec in doc.sections:
        sec.page_width=Cm(21); sec.page_height=Cm(29.7)
        sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.0)

    # ---------- Cover ----------
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; _sp(p,80,4)
    _run(p,"GEOモニタリングレポート",bold=True,size=28,color=C_PRIMARY)
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; _sp(p,4,20)
    _run(p, period, bold=True, size=22, color=C_ACCENT)
    p = doc.add_paragraph(); _sp(p,0,20)
    pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
    bot=OxmlElement("w:bottom"); bot.set(qn("w:val"),"single")
    bot.set(qn("w:sz"),"8"); bot.set(qn("w:color"),HX_ACCENT)
    pBdr.append(bot); pPr.append(pBdr)
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; _sp(p,20,4)
    _run(p,"Dears Wedding（運営：Arluis）",size=13,color=C_GRAY)
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; _sp(p,4,0)
    _run(p,"AI言及率・誤情報モニタリング",size=11,color=C_GRAY)
    doc.add_page_break()

    # ---------- Executive Summary ----------
    _h1(doc,"エグゼクティブサマリー")
    for line in [f"調査期間：{period}",
                 "対象AIプラットフォーム：ChatGPT・Gemini・Perplexity・AI Overview（4社）"]:
        p=doc.add_paragraph(); _sp(p,3,3); _run(p,line,size=11)
    doc.add_paragraph()
    p=doc.add_paragraph(); _sp(p,3,3)
    _run(p,f"総合言及率：{overall}（{total_m} / {total_q} クエリ）",bold=True,size=11)
    p=doc.add_paragraph(); _sp(p,3,12)
    _run(p,f"誤情報検出件数：{n_mis}件",bold=True,size=11,
         color=C_RED if n_mis>0 else RGBColor(0x33,0x33,0x33))

    # KPI cards (2-row table)
    t=doc.add_table(rows=2,cols=4); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,pk in enumerate(PROVIDERS):
        c=t.rows[0].cells[i]; _bg(c,HX_PRIMARY); _borders(c,HX_ACCENT)
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _run(pp,LABELS[pk],bold=True,size=10,color=C_WHITE)
    for i,pk in enumerate(PROVIDERS):
        c=t.rows[1].cells[i]; _bg(c,HX_LIGHT); _borders(c,HX_ACCENT)
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        st=mr.get(pk,{}); rate=f"{st.get('mention_rate',0)*100:.1f}%"
        pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _run(pp,rate,bold=True,size=16,color=C_ACCENT)
        p2=c.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _run(p2,f"{st.get('total',0)}クエリ",size=9,color=C_GRAY)
    doc.add_paragraph()

    # ---------- Detail table ----------
    _h1(doc,"1. AI別 言及率 詳細")
    p=doc.add_paragraph(); _sp(p,3,8)
    _run(p,"各AIプラットフォームにおける Dears Wedding（Arluis）の言及状況です。",size=11)
    dt=doc.add_table(rows=1,cols=4); dt.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(["プロバイダー","言及率",
                           "言及クエリ数","総クエリ数"]):
        c=dt.rows[0].cells[i]; _bg(c,HX_PRIMARY); _borders(c,HX_PRIMARY)
        pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _run(pp,h,bold=True,size=10,color=C_WHITE)
    for idx,pk in enumerate(PROVIDERS):
        rc=dt.add_row().cells; bg=HX_WHITE if idx%2==0 else HX_LGRAY
        st=mr.get(pk,{})
        vals=[LABELS[pk],f"{st.get('mention_rate',0)*100:.1f}%",
              str(st.get("mentioned",0)),str(st.get("total",0))]
        for i,v in enumerate(vals):
            _bg(rc[i],bg); _borders(rc[i])
            pp=rc[i].paragraphs[0]
            pp.alignment=WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER
            _run(pp,v,size=10)
    doc.add_paragraph()

    # ---------- Misinformation ----------
    _h1(doc,"2. 誤情報アラート")
    p=doc.add_paragraph(); _sp(p,3,8)
    _run(p,f"今月は {n_mis}件 の誤情報が検出されました。"
           +("早急な対策が推奨されます。" if n_mis>0 else ""),size=11)
    if n_mis>0:
        mt=doc.add_table(rows=1,cols=3); mt.alignment=WD_TABLE_ALIGNMENT.CENTER
        for i,h in enumerate(["AI","検出内容（要約）","クエリ"]):
            c=mt.rows[0].cells[i]; _bg(c,HX_RED); _borders(c,HX_RED)
            pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            _run(pp,h,bold=True,size=10,color=C_WHITE)
        for idx,r in enumerate(mis[:20]):
            rc=mt.add_row().cells; bg=HX_WHITE if idx%2==0 else HX_PINK
            det=(r.get("misinformation_detail") or "")[:150]
            if len(r.get("misinformation_detail",""))>150: det+="…"
            qry=(r.get("prompt_text") or "")[:40]
            if len(r.get("prompt_text",""))>40: qry+="…"
            for i,v in enumerate([LABELS.get(r.get("ai_provider",""),""),det,qry]):
                _bg(rc[i],bg); _borders(rc[i],HX_REDLT)
                _run(rc[i].paragraphs[0],v,size=9)
    else:
        p=doc.add_paragraph(); _sp(p,3,8)
        _run(p,"今月の誤情報検出はありませんでした。",size=11,color=C_GREEN)
    doc.add_paragraph()

    # ---------- Recommendations ----------
    _h1(doc,"3. 所見と推奨アクション")
    _h2(doc,"現状分析")
    for b in [
        f"総合言及率 {overall} は、ブランド認知がAI回答に反映されている割合を示します。",
        "誤情報として最も多いパターン：「東京都内専門・少人析70名以下」に反する記述（全国展開・大規模月次対応など）。",
        "Gemini・Perplexityでは他社ブランドと混同される傾向が見られます。",
    ]:
        p=doc.add_paragraph(style="List Bullet"); _sp(p,2,2); _run(p,b,size=11)
    _h2(doc,"推奨アクション")
    for a in [
        "《高優先度》公式サイトのコンテンツを「東京都内・少人析70名以下専門」と明示するページを強化し、AIクローラーに正確な情報を提供する。",
        "《中優先度》誤情報が多いクエリカテゴリ（エリア・規模）に対し、Q&A形式のFAQページを追加する。",
        "《継続モニタリング》来月も同クエリセットで再計測し、言及率の推移をトラッキングする。",
    ]:
        p=doc.add_paragraph(style="List Bullet"); _sp(p,2,2); _run(p,a,size=11)
    doc.add_paragraph()

    # ---------- Data source ----------
    _h1(doc,"4. データソース")
    p=doc.add_paragraph(); _sp(p,3,4)
    _run(p,"詳細データはGoogle Sheetsで確認できます：",size=11)
    p=doc.add_paragraph(); _sp(p,3,12)
    if sh_url: _hyperlink(p, sh_url, sh_url)
    else: _run(p,"(スプレッドシートURL未設定)",size=11,color=C_GRAY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; _sp(p,20,0)
    _run(p,"— 以上 —",size=10,color=C_GRAY)

    doc.save(str(out_path))
    logger.info("Word saved: %s", out_path)
    return out_path


def fetch_data(db_url, client_id, year, month, sheet_id) -> dict:
    saver = MentionSaver(db_url)
    mr    = saver.fetch_mention_rate(client_id, year, month)
    rows  = saver.fetch_recent_analysis(client_id, limit=500)
    pfx   = f"{year}-{month:02d}"
    month_rows = [r for r in rows
                  if str(r.get("responded_at","")).startswith(pfx)
                  or str(r.get("created_at","")).startswith(pfx)] or rows
    mis   = [r for r in month_rows if r.get("misinformation_flag")]
    def clean(s): return {"mention_rate":float(s.get("mention_rate",0)),
                          "total":int(s.get("total",0)),"mentioned":int(s.get("mentioned",0))}
    return {
        "year":year,"month":month,
        "mention_rate":{k:clean(v) for k,v in mr.items()},
        "misinfo_rows":[{"ai_provider":str(r.get("ai_provider","")),"prompt_text":str(r.get("prompt_text","")),
                         "misinformation_flag":bool(r.get("misinformation_flag",False)),
                         "misinformation_detail":str(r.get("misinformation_detail",""))} for r in mis],
        "spreadsheet_url":f"https://docs.google.com/spreadsheets/d/{sheet_id}" if sheet_id else "",
    }


def to_pdf(docx_path: Path):
    lo_candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    lo = next((p for p in lo_candidates if Path(p).exists()), None)
    if lo is None:
        logger.warning("LibreOffice が見つかりません。PDF 変換をスキップします。")
        logger.warning("PDF が必要な場合は https://www.libreoffice.org からインストールしてください。")
        return None
    try:
        r = subprocess.run([lo,"--headless","--convert-to","pdf",
                            "--outdir",str(docx_path.parent),str(docx_path)],
                           capture_output=True, text=True)
        if r.returncode!=0:
            logger.warning("PDF 変換失敗。Word ファイルは生成済みです。")
            return None
        return docx_path.with_suffix(".pdf")
    except Exception as e:
        logger.warning("PDF 変換エラー: %s", e)
        return None


def main():
    now = datetime.now(timezone.utc)
    p   = argparse.ArgumentParser()
    p.add_argument("--year",  type=int, default=now.year)
    p.add_argument("--month", type=int, default=now.month)
    p.add_argument("--client-id", default=DEFAULT_CLIENT_ID)
    p.add_argument("--no-pdf", action="store_true")
    args = p.parse_args()

    db_url = os.environ.get("DATABASE_URL")
    if not db_url:
        print("ERROR: DATABASE_URL not set"); sys.exit(1)

    REPORTS_DIR.mkdir(exist_ok=True)
    out = REPORTS_DIR / f"GEO_Report_{args.year}_{args.month:02d}.docx"
    data = fetch_data(db_url, args.client_id, args.year, args.month,
                      os.environ.get("GEO_SPREADSHEET_ID"))
    build_report(data, out)
    pdf = None if args.no_pdf else to_pdf(out)

    print(f"\n{'='*50}\n  Word: {out}")
    if pdf: print(f"  PDF : {pdf}")
    print('='*50)

if __name__=="__main__":
    main()
